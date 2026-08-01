import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    """Set the background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def add_page_number(run):
    """Inserts a page number field inside a run."""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def create_document():
    doc = docx.Document()
    
    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
        # Add footer with page number
        footer = section.footer
        footer_p = footer.paragraphs[0]
        footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer_run = footer_p.add_run("Page ")
        footer_run.font.name = 'Calibri'
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(128, 128, 128)
        add_page_number(footer_run)
        
    # Styles Setup
    styles = doc.styles
    
    # Custom colors
    NAVY = RGBColor(30, 58, 138)       # #1E3A8A - Primary Headers
    SLATE = RGBColor(71, 85, 105)      # #475569 - Secondary Headers
    CHARCOAL = RGBColor(30, 41, 59)    # #1E293B - Body text
    RED = RGBColor(185, 28, 28)        # #B91C1C - Highlights/Warnings
    
    # ------------------ COVER PAGE ------------------
    for _ in range(3):
        doc.add_paragraph()
        
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("TECHNICAL ARCHITECTURE & SYSTEM DOCUMENTATION")
    title_run.bold = True
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = NAVY
    title_run.font.name = 'Segoe UI'
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_p.add_run("Agent-Stream Multi-Tenant Voice AI Platform\nEnd-to-End Enterprise SIP Telephony Integrations")
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = SLATE
    sub_run.font.name = 'Segoe UI'
    
    for _ in range(5):
        doc.add_paragraph()
        
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta_p.add_run("Prepared for: Chauwk Enterprise Solutions\nPrepared by: Advanced AI Engineering Team\nDate: August 2026\nVersion: 2.0.0")
    meta_run.font.size = Pt(11)
    meta_run.font.color.rgb = CHARCOAL
    meta_run.font.name = 'Calibri'
    
    doc.add_page_break()
    
    # Helper to add stylized headings
    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(18)
        run.font.color.rgb = NAVY
        run.font.name = 'Segoe UI'
        return p

    def add_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = SLATE
        run.font.name = 'Segoe UI'
        return p

    def add_heading_3(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(11.5)
        run.font.color.rgb = CHARCOAL
        run.font.name = 'Segoe UI'
        return p

    def add_body_text(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.color.rgb = CHARCOAL
        run.font.name = 'Calibri'
        return p

    def add_bullet(text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.color.rgb = CHARCOAL
        run.font.name = 'Calibri'
        return p

    def add_code_block(text):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.4)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.0
        
        # Add border style or simple background shading if possible, or print code
        run = p.add_run(text)
        run.font.size = Pt(9.5)
        run.font.name = 'Consolas'
        run.font.color.rgb = RGBColor(30, 41, 59)
        return p

    def add_warning(text):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run("⚠️ WARNING: " + text)
        run.font.size = Pt(11)
        run.font.color.rgb = RED
        run.font.name = 'Calibri'
        run.bold = True
        return p

    # ------------------ SECTION 1 ------------------
    add_heading_1("1. Executive Overview & Connection Architecture")
    add_body_text(
        "The Agent-Stream Voice AI Bot Platform is a high-performance, multi-tenant enterprise system "
        "designed to handle inbound and outbound voice calls with low latency. The platform acts as a bridge "
        "between Exotel's telephony network and advanced AI engines, enabling natural, human-like voice conversations."
    )
    
    add_heading_2("1.1 Network Connection Flow: Exotel to AWS EC2")
    add_body_text(
        "The entire platform is hosted within a secure AWS EC2 Instance located in the Mumbai region (AP-SOUTH-1). "
        "The connection architecture supports direct audio streaming over SIP trunks and WebSockets. The connection path "
        "comprises the following stages:"
    )
    add_bullet("Exotel Telephony Network: Telephony endpoints and virtual numbers (DIDs) receive incoming calls from customers or dial out leg-by-leg connections.")
    add_bullet("AWS Security Group & VPC: Exposes TCP/UDP port 5060 (SIP Signaling), UDP ports 10000-20000 (RTP Audio streaming), and TCP port 5000/5002 (FastAPI gateway and Web Dashboard) only to whitelisted IP blocks.")
    add_bullet("PJSIP C-Library SIP Server: Runs natively inside the dockerized container to terminate carrier SIP connections and capture raw RTP packets directly.")
    add_bullet("FastAPI Web Server: Processes admin API routes, logs, multi-tenant uploads, and handles browser WebSocket streaming connection legs.")

    # ------------------ SECTION 2 ------------------
    add_heading_1("2. Telephony & Connection Integration Flow")
    
    add_heading_2("2.1 Inbound SIP Trunking (vSIP)")
    add_body_text(
        "Inbound voice calls connect directly to our self-hosted PJSIP stack. This design completely eliminates "
        "the need for Exotel Applet URLs or visual call flow setups, reducing connection overhead and latency. "
        "The inbound workflow follows these steps:"
    )
    add_bullet("INVITE: The Exotel carrier trunk sends a SIP INVITE request targeting our public IP address at port 5060.")
    add_bullet("Security Whitelisting: PJSIP validates the calling address against a strict set of authentic carrier domains (*.exotel.in, *.exotel.com) and IP prefixes.")
    add_bullet("Answer: The server replies with a 200 OK, initiating the media session.")
    add_bullet("RTP Audio Bridging: PJSIP initializes an audio media stream mapping carrier G.711 / PCM audio into memory, making it immediately available for AI engines.")

    add_heading_2("2.2 Outbound REST Connect API")
    add_body_text(
        "Outbound campaigns and manual dialers leverage Exotel's peer-to-peer bridging API (`/v1/Accounts/{AccountSid}/Calls/connect.json`). "
        "Instead of dialing an outbound SIP client, the REST API creates a bridged call:"
    )
    add_bullet("Leg A (Customer Dialing): Exotel dials the customer's phone number ('From' parameter).")
    add_bullet("Leg B (Agent Bridging): As soon as the customer picks up, Exotel dials our virtual DID number ('To' parameter).")
    add_bullet("Trunk Inbound routing: The virtual DID maps directly to our SIP trunk, launching the server's AI bot immediately as a standard inbound leg.")
    add_body_text(
        "Important Parameter Sanitation: The system cleans and formats the destination number by removing "
        "country codes (+91) and prefixing a leading '0' (e.g. 09553856533). Exotel's gateway strictly requires "
        "this format for Indian mobile numbers."
    )

    add_heading_2("2.3 Browser Audio Streaming & Web Component")
    add_body_text(
        "For browser-based voice agents, the application includes a custom Web Component (<agent-stream-voice>) "
        "backed by `/static/widget.js`. It establishes a secure WebSocket connection directly to the FastAPI server. "
        "To ensure compatibility with Deepgram, client-side microphone input is resampled to 16kHz before transmission, "
        "and synthesized TTS audio is received at 16kHz to prevent static/clicking noise."
    )

    # ------------------ SECTION 3 ------------------
    add_heading_1("3. Dual AI Processing Engine Modes")
    add_body_text(
        "The system runs in one of two modes, selectable via the VOICE_BOT_MODE environment variable."
    )

    add_heading_2("3.1 Mode 1: Modular Pipeline (Deepgram + Gemini + Sarvam)")
    add_body_text(
        "The modular mode decouples Speech-to-Text, Language Modeling, and Text-to-Speech into separate specialized APIs. "
        "This offers granular control over voice tuning and is optimized for Indian vernacular languages:"
    )
    add_bullet("STT (Deepgram): Captures user audio from the telephony RTP stream and streams it to the Deepgram WebSocket. It has been upgraded to the high-performance 'nova-3' model. Dynamic keyword boosting is used to prioritize custom product names.")
    add_bullet("LLM (Gemini 2.5 Flash): Receives text transcripts. It connects to the Google GenAI SDK (Vertex AI/AI Studio) to handle RAG queries, email tool calls, and call terminations. Context injection is optimized by fetching only the top-2 matching chunks to minimize token usage.")
    add_bullet("TTS (Sarvam bulbul:v3): Text responses are sent to Sarvam's WebSocket. If the WebSocket connection is unstable or silent, the bot automatically falls back to Sarvam's HTTP REST API. Supports hot-reloaded speakers (shubh, neha, ishita), speaking speed, and volume gain.")
    add_bullet("Silence Monitoring & VAD: Employs a 3-intimation silence disconnect policy. If the user is silent for 8 seconds, the system injects a polite reminder (e.g., 'Are you still there?'). If silence continues after three reminders, the call is disconnected automatically to save credits.")

    add_heading_2("3.2 Mode 2: OpenAI Realtime Speech-to-Speech")
    add_body_text(
        "The realtime mode utilizes a direct bidirectional WebSocket connection to OpenAI's Realtime API."
        "It supports native speech-to-speech without intermediate text conversion, providing extremely low latency."
        "To eliminate clicking noises, a stateful resampling buffer was implemented to smooth out audio packet boundaries "
        "when converting between telephony formats (G.711/16kHz PCM) and OpenAI's native audio rate."
    )

    # ------------------ SECTION 4 ------------------
    add_heading_1("4. Codebase Directory & File-by-File Analysis")
    
    files_data = [
        ("api_gateway.py", "Main gateway file running FastAPI. Hosts administrative endpoints, serves dashboard assets, manages company/agent metadata, and routes browser WebSockets."),
        ("config.py", "Centralized configuration loader. Uses pydantic-style variables to load .env configs, validate keys (OpenAI, Gemini, Sarvam, Deepgram), and handles fallback checks."),
        ("core/sip_server.py", "Native PJSIP python bindings wrapper. Starts the SIP stack, handles incoming carrier INVITE requests, enforces whitelisting, and bridges telephony calls to the AI bots."),
        ("core/modular_sales_bot.py", "The core driver for modular mode. Manages Websocket connections to Deepgram STT, executes Gemini LLM loops with system instructions, calls RAG context queries, and feeds text responses into Sarvam TTS."),
        ("core/openai_realtime_sales_bot.py", "The core driver for OpenAI Realtime mode. Manages the bidirectional WebSocket connection, configures tool definitions, handles audio resampling, and manages barge-in interruptions."),
        ("core/email_client.py", "Asynchronous SMTP email client. Uses Python's standard smtplib wrapped in asyncio.to_thread to send non-blocking email alerts via Zoho SMTP (smtp.zoho.in:465)."),
        ("core/rag_manager.py", "Retrieval-Augmented Generation manager. Integrates with ChromaDB, extracts text from uploaded PDFs/Word files, and handles semantic similarity vector searches."),
        ("core/analytics_manager.py", "Generates call insights. On call completion, extracts customer name, email, phone, sentiment, and meeting consent from the transcript using Gemini and saves the results in MongoDB Atlas."),
        ("core/agent_resolver.py", "Resolves active agent profiles by phoneNumber, agent_id, or company_id to dynamically map prompts and greetings during incoming call routing."),
        ("core/bot_framework.py", "Dynamic bot framework managing templates (sales, support, debt recovery) with filesystem-watch configurations for live hot-reloading."),
        ("core/bot_launcher.py", "FastAPI startup task wrapper that initializes global servers, loads pre-cached agent greetings, and monitors background tasks."),
        ("core/mongo_manager.py", "Persistent database connector. Establishes connection to MongoDB Atlas for indexing agents, enterprise companies, campaigns, documents, and call logs.")
    ]

    for filename, description in files_data:
        add_heading_3(filename)
        add_body_text(description)

    # ------------------ SECTION 5 ------------------
    add_heading_1("5. Detailed Environment Configuration (.env)")
    add_body_text("The application relies on the following key environment variables in the master .env file:")
    
    env_keys = [
        ("VOICE_BOT_MODE", "Sets the active AI engine ('modular' or 'realtime')."),
        ("DISABLE_AI_ENGINES", "Emergency safety switch. When true, bypasses API validations and rejects all incoming calls with 503 Service Unavailable."),
        ("DEEPGRAM_API_KEY", "Authentication token for Deepgram Speech-to-Text streaming."),
        ("GEMINI_API_KEY", "Authentication token for Google Gemini LLM API (AI Studio)."),
        ("SARVAM_API_KEY", "Authentication token for Sarvam AI Text-to-Speech API."),
        ("OPENAI_API_KEY", "Authentication token for OpenAI Realtime API."),
        ("SIP_PUBLIC_IP", "The public IP of your EC2 instance (e.g. 3.111.29.229) used for SIP trunk routing."),
        ("INBOUND_SIP_ENABLED", "Enables or disables receiving incoming SIP calls."),
        ("EXOTEL_ACCOUNT_SID", "Exotel Account SID username for REST calls."),
        ("EXOTEL_API_KEY / TOKEN", "Credentials for making Exotel API requests."),
        ("EXOTEL_FROM_NUMBER", "Your virtual DID phone number (e.g. 04040377112)."),
        ("SMTP_HOST / USER / PASSWORD", "Credentials for sending Zoho emails (smtp.zoho.in, port 465)."),
        ("DB_URL", "MongoDB Atlas connection string."),
        ("AWS_S3_BUCKET_NAME", "Target AWS S3 bucket name for RAG document sync.")
    ]

    for key, purpose in env_keys:
        add_bullet(f"{key}: {purpose}")

    # ------------------ SECTION 6 ------------------
    doc.add_page_break()
    add_heading_1("6. FastAPI REST & WebSocket APIs Reference")
    
    add_heading_2("6.1 Admin & Management REST APIs")
    add_bullet("POST /api/v1/bot/config: Hot-reloads active voice settings (speaker name, speed/pace, audio gain).")
    add_bullet("POST /companies/{id}/documents: Uploads PDF, DOCX, TXT, or XLSX templates for RAG indexing.")
    add_bullet("GET /api/v1/calls/campaigns: Retrieves outbound call campaigns by enterprise_id and agent_id.")
    add_bullet("GET /admin/all: Admin endpoint to list all configured agents across all enterprises.")
    add_bullet("DELETE /exotel-deletekb-items/{doc_id}: Deletes parsed RAG document items.")
    
    add_heading_2("6.2 Web Simulator & WebSocket APIs")
    add_bullet("WS /api/v1/ws/browser-stream: Bidirectional browser WebSocket streaming for voice web-widgets.")
    add_bullet("POST /simulate/history: Saves simulated conversation logs to simulated_conversations collection.")

    # ------------------ SECTION 7 ------------------
    add_heading_1("7. Git Commit History & Project Milestones")
    add_body_text("Below is a chronological breakdown of major development milestones resolved through GitHub commits:")
    
    commits = [
        ("3588558", "Integrated phoneNumber field into the core agent creation document schemas to enable direct lookup."),
        ("e2c0fd0", "Upgraded default Deepgram model to nova-3-phonecall in .env configuration for improved Indian languages word accuracy."),
        ("bb7581a", "Swapped Deepgram WebSocket params to use keyterm= instead of keywords= to prevent HTTP 400 rejection on keywords boosting."),
        ("a536cf0", "Fixed greeting latency. Implemented async deduplication to eliminate delay when connecting STT streams (Telugu, Hindi, Tamil)."),
        ("87f15e7", "Implemented Unicode script-based sentence parser to automatically detect script type (e.g. Devanagari) for per-sentence TTS code-switching."),
        ("21c2e84", "Refactored keyword boosting in Deepgram to be dynamically loaded per agent from their metadata rather than hardcoded."),
        ("78b390c", "Fixed session_state key bug where live calls were skipping RAG knowledge context fetching due to incorrect attribute name check."),
        ("98f5496", "Fixed playback silence issues. Implemented automatic fallback routing to Sarvam REST HTTP API with gain whenever WebSocket drops."),
        ("9308c81", "Fixed NameError re and session_from_phone resolution bugs during incoming call setup in modular sales bot."),
        ("5bdc811", "Implemented inbound customer callback routing. Automatically routes inbound callers to the exact agent who dialed them during outbound campaigns."),
        ("4737754", "Latency Optimization: Implemented dynamic per-turn top-2 RAG context injection to minimize token processing and payload latency."),
        ("056d944", "Bulk Calling Campaigns: Integrated Excel file (.xlsx) parser to support uploading massive lead campaign templates on the dashboard."),
        ("c7065b4", "Documentation: Added complete system architecture flows and comparisons in the primary repository README.md.")
    ]

    for sha, detail in commits:
        add_bullet(f"Commit {sha}: {detail}")

    # ------------------ SECTION 8 ------------------
    doc.add_page_break()
    add_heading_1("8. Lessons Learned & Mistakes to Avoid")
    
    add_warning("Private IP Subnet Whitelisting in Production")
    add_body_text(
        "Mistake: Whitelisting all local private network IPs (10.x.x.x, 172.16.x.x) for developer convenience. "
        "Impact: SIP scanners spoofed their SIP headers to match these subnets and bypassed authentication. This allowed them to connect and consume Sarvam AI credits on empty calls. "
        "Prevention: Disable private IP whitelisting in production when Config.SIP_PUBLIC_IP is configured."
    )
    
    add_warning("Missing Seeding & System Filters in Email Alerts")
    add_body_text(
        "Mistake: Triggering emails on any transcript history containing user roles. "
        "Impact: Automated greeting prompts and system silence messages ('System: Customer silent...') were saved with the user role. This triggered continuous email alerts on empty/scanner calls. "
        "Prevention: Explicitly filter out any text starting with 'System:' or matching the seeding greeting from the user interaction check."
    )
    
    add_warning("PJSIP Garbage Collection Assertion Crashes")
    add_body_text(
        "Mistake: Failing to pin active PJSIP call objects in Python scope. "
        "Impact: Python's garbage collector destroyed call objects mid-session, causing random crashes with 'libRegisterThread' assertion errors. "
        "Prevention: Explicitly register and pin all active PJSIP call objects in a global thread-safe map, and unregister them only during clean hangup events."
    )

    add_warning("Exotel Outbound Number Formatting")
    add_body_text(
        "Mistake: Passing raw international format numbers (+91) to the Exotel Connect API. "
        "Impact: Exotel rejected calls with validation errors. "
        "Prevention: Format Indian mobile numbers by removing '+91' and adding a leading '0' prefix."
    )

    # ------------------ SECTION 9 ------------------
    add_heading_1("9. Future Enhancements & System Roadmap")
    add_bullet("Restricting Port 5060 in AWS Security Group: Limit inbound port 5060 traffic only to Exotel's IP ranges at the network firewall level.")
    add_bullet("Auto-Restart Watchdog: Implement a supervisord daemon or shell watchdog script to monitor PJSIP process memory and automatically restart the service if it crashes.")
    add_bullet("Semantic Caching: Implement Redis semantic caching for RAG queries to return immediate cached answers for recurring customer questions, reducing LLM costs.")
    add_bullet("Dynamic Voice Tuning Dashboard: Expand the web UI to support configuring distinct greeting prompts and allowed language lists per agent dynamically.")

    # Save document
    output_path = "c:\\Users\\premm\\Desktop\\Repos\\Agent-Stream-Voice-Agent\\Voice_AI_Bot_Project_Documentation.docx"
    doc.save(output_path)
    print(f"Technical Documentation Document created successfully at {output_path}!")

if __name__ == "__main__":
    create_document()
